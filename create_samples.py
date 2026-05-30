from docx import Document
import os

os.makedirs("sample_resumes", exist_ok=True)

# Resume 1 - Strong match
doc1 = Document()
doc1.add_heading("Alice Johnson", 0)
doc1.add_paragraph("Email: alice@email.com | Phone: 9876543210")
doc1.add_heading("Skills", level=1)
doc1.add_paragraph("Python, scikit-learn, pandas, numpy, TensorFlow, NLP, machine learning, data science")
doc1.add_heading("Experience", level=1)
doc1.add_paragraph("Data Scientist at TechCorp (2 years)\n- Built ML models using scikit-learn and TensorFlow\n- Worked on NLP text classification projects\n- Used pandas and numpy for data processing")
doc1.add_heading("Education", level=1)
doc1.add_paragraph("B.Tech Computer Science - IIT Delhi (2022)")
doc1.save("sample_resumes/alice_johnson.docx")

# Resume 2 - Medium match
doc2 = Document()
doc2.add_heading("Bob Smith", 0)
doc2.add_paragraph("Email: bob@email.com | Phone: 9876543211")
doc2.add_heading("Skills", level=1)
doc2.add_paragraph("Python, Django, REST APIs, PostgreSQL, JavaScript, HTML, CSS")
doc2.add_heading("Experience", level=1)
doc2.add_paragraph("Backend Developer at WebSolutions (3 years)\n- Built REST APIs using Django\n- Managed PostgreSQL databases\n- Some experience with Python scripting")
doc2.add_heading("Education", level=1)
doc2.add_paragraph("B.Tech Computer Science - NIT Calicut (2021)")
doc2.save("sample_resumes/bob_smith.docx")

# Resume 3 - Weak match
doc3 = Document()
doc3.add_heading("Carol Davis", 0)
doc3.add_paragraph("Email: carol@email.com | Phone: 9876543212")
doc3.add_heading("Skills", level=1)
doc3.add_paragraph("Photoshop, Illustrator, Figma, UI/UX Design, Canva, Sketch")
doc3.add_heading("Experience", level=1)
doc3.add_paragraph("UI/UX Designer at DesignStudio (2 years)\n- Designed mobile app interfaces\n- Created brand identity materials\n- Conducted user research")
doc3.add_heading("Education", level=1)
doc3.add_paragraph("B.Des Visual Communication - NID Ahmedabad (2022)")
doc3.save("sample_resumes/carol_davis.docx")

print("✅ 3 sample resumes created in sample_resumes/ folder!")